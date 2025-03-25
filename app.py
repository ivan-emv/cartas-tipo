import streamlit as st
from docx import Document
from io import BytesIO
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def replace_text_in_docx(template_path, replacements):
    doc = Document(template_path)
    
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        
                        if key in ["(INSERTETRAMO)", "(MODODETRANSPORTE)", "(FECHA1)"]:
                            run.font.name = "Arial Black"
                            run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    for para in doc.paragraphs:
        if any(tag in para.text for tag in ["(DATOS1)", "(DATOS2)", "(DATOS3)"]):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    return doc

st.title("Generador de Documentos de Transporte")

# 🔧 Ocultar la barra superior y el menú de Streamlit
hide_streamlit_style = """
    <style>
        #MainMenu {visibility: hidden;}
        header {visibility: hidden;}
        footer {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

idioma = st.selectbox("Seleccione el idioma del documento", ["Español", "Portugués", "Inglés"])

if idioma == "Español":
    template_path = "Carta Tipo - Transporte.docx"
    opciones_transporte = {"AVIÓN": "AVIÓN", "TREN": "TREN", "AUTOBÚS": "AUTOBÚS"}
elif idioma == "Portugués":
    template_path = "Carta Tipo - Transporte - POR.docx"
    opciones_transporte = {"AVIÓN": "AVIÃO", "TREN": "TREM", "AUTOBÚS": "ÔNIBUS"}
elif idioma == "Inglés":
    template_path = "Carta Tipo - Transporte - ENG.docx"
    opciones_transporte = {"AVIÓN": "PLANE", "TREN": "TRAIN", "AUTOBÚS": "BUS"}

txt_nombre = st.text_input("Ingrese Nombre de PAX")
txt_localizador = st.text_input("Ingrese Localizador(es)")
txt_tramo = st.text_input("Ingrese Tramo")
opcion_transporte = st.selectbox("Seleccione Modo de Transporte", ["AVIÓN", "TREN", "AUTOBÚS"])
txt_fecha = st.text_input("Ingrese Fecha")
txt_datos1 = st.text_area("Ingrese Datos 1")
txt_datos2 = st.text_area("Ingrese Datos 2")
txt_datos3 = st.text_area("Ingrese Datos 3")

modo_traducido = opciones_transporte[opcion_transporte]

replacements = {
    "(INSERTENOMBRE)": txt_nombre,
    "(LOCALIZADOR)": txt_localizador,
    "(INSERTETRAMO)": txt_tramo,
    "(MODODETRANSPORTE)": modo_traducido,
    "(FECHA1)": txt_fecha,
    "(DATOS1)": txt_datos1,
    "(DATOS2)": txt_datos2,
    "(DATOS3)": txt_datos3
}

if st.button("Generar Documento"):
    doc = replace_text_in_docx(template_path, replacements)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    file_name = f"{txt_localizador}.docx"

    st.download_button(label="Descargar Documento", data=buffer, file_name=file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
